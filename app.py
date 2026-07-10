"""
НейроПост v2.0 — Графический интерфейс
Запуск: python app.py
"""
import json
import os
import sys
import threading
import subprocess
from datetime import date
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog

# ── Пути ──────────────────────────────────────────────────────────────────────
BASE_DIR     = os.path.dirname(os.path.abspath(__file__))
PROFILES_DIR = os.path.join(BASE_DIR, "profiles")
OWNER_FILE   = os.path.join(PROFILES_DIR, "owner.json")
OUTPUT_POSTS = os.path.join(BASE_DIR, "output", "posts")
OUTPUT_AUDIT = os.path.join(BASE_DIR, "output", "audits")

# ── Цвета ─────────────────────────────────────────────────────────────────────
BG       = "#F5F6FA"
SIDEBAR  = "#2C3E50"
ACCENT   = "#27AE60"
ACCENT2  = "#2980B9"
WHITE    = "#FFFFFF"
TEXT     = "#2C3E50"
GRAY     = "#95A5A6"
LIGHT    = "#ECF0F1"
RED      = "#E74C3C"
ORANGE   = "#F39C12"

FONT_H1   = ("Arial", 20, "bold")
FONT_H2   = ("Arial", 14, "bold")
FONT_BODY = ("Arial", 11)
FONT_BTN  = ("Arial", 11, "bold")
FONT_SMALL= ("Arial", 9)


# ══════════════════════════════════════════════════════════════════════════════
#  Вспомогательные функции
# ══════════════════════════════════════════════════════════════════════════════

def load_json(path):
    with open(path, encoding="utf-8") as f:
        return json.load(f)

def save_json(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def get_profiles():
    os.makedirs(PROFILES_DIR, exist_ok=True)
    result = []
    for fname in sorted(os.listdir(PROFILES_DIR)):
        if fname.endswith(".json"):
            try:
                data = load_json(os.path.join(PROFILES_DIR, fname))
                result.append((fname, data))
            except Exception:
                pass
    result.sort(key=lambda x: 0 if x[1].get("_type") == "owner" else 1)
    return result

def name_to_filename(name):
    keep = "abcdefghijklmnopqrstuvwxyz0123456789_абвгдеёжзийклмнопрстуфхцчшщъыьэюя"
    safe = "".join(c for c in name.lower().replace(" ", "_") if c in keep)
    return safe or "client"

def run_in_thread(fn, *args):
    t = threading.Thread(target=fn, args=args, daemon=True)
    t.start()


# ══════════════════════════════════════════════════════════════════════════════
#  Главное окно
# ══════════════════════════════════════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("НейроПост v2.0")
        self.geometry("1100x720")
        self.minsize(900, 600)
        self.configure(bg=BG)
        self.resizable(True, True)

        # Иконка (буква Н)
        try:
            self.iconbitmap(default="")
        except Exception:
            pass

        self.active_profile_file = tk.StringVar(value="")
        self.active_profile_data = {}

        self._build_layout()
        self._check_first_run()

    # ── Построение макета ──────────────────────────────────────────────────
    def _build_layout(self):
        # Боковая панель
        self.sidebar = tk.Frame(self, bg=SIDEBAR, width=220)
        self.sidebar.pack(side="left", fill="y")
        self.sidebar.pack_propagate(False)

        # Логотип
        tk.Label(self.sidebar, text="НейроПост", font=("Arial", 18, "bold"),
                 bg=SIDEBAR, fg=WHITE).pack(pady=(28, 2))
        tk.Label(self.sidebar, text="v2.0", font=FONT_SMALL,
                 bg=SIDEBAR, fg=GRAY).pack(pady=(0, 20))

        ttk.Separator(self.sidebar, orient="horizontal").pack(fill="x", padx=20, pady=5)

        # Профиль (кнопка выбора)
        tk.Label(self.sidebar, text="Активный профиль:", font=FONT_SMALL,
                 bg=SIDEBAR, fg=GRAY).pack(pady=(15, 2))
        self.profile_btn = tk.Button(
            self.sidebar, text="— не выбран —",
            font=("Arial", 10, "bold"), bg=ACCENT, fg=WHITE,
            relief="flat", cursor="hand2", wraplength=180,
            command=self._open_profile_selector,
        )
        self.profile_btn.pack(padx=15, fill="x", pady=(0, 15))

        ttk.Separator(self.sidebar, orient="horizontal").pack(fill="x", padx=20, pady=5)

        # Кнопки меню
        menu_items = [
            ("✍️  Написать пост",           self._show_generator),
            ("📅  Контент-план (Excel)",     self._show_calendar),
            ("📢  Опубликовать в Telegram",  self._show_publish_tg),
            ("📣  Опубликовать в VK",        self._show_publish_vk),
            ("🔍  НейроАудит клиента",       self._show_audit),
        ]
        for label, cmd in menu_items:
            self._sidebar_btn(label, cmd)

        ttk.Separator(self.sidebar, orient="horizontal").pack(fill="x", padx=20, pady=10)

        self._sidebar_btn("👥  Управление профилями", self._show_profiles, color=ACCENT2)

        # Версия внизу
        tk.Label(self.sidebar, text="© 2026 НейроПост", font=FONT_SMALL,
                 bg=SIDEBAR, fg=GRAY).pack(side="bottom", pady=10)

        # Контентная область
        self.content = tk.Frame(self, bg=BG)
        self.content.pack(side="left", fill="both", expand=True)

        self._show_home()

    def _sidebar_btn(self, text, cmd, color=None):
        c = color or "#3D5166"
        b = tk.Button(
            self.sidebar, text=text, font=("Arial", 10),
            bg=c, fg=WHITE, relief="flat", anchor="w",
            padx=18, pady=10, cursor="hand2",
            activebackground=ACCENT, activeforeground=WHITE,
            command=cmd,
        )
        b.pack(fill="x", padx=10, pady=2)
        return b

    # ── Очистка контента ───────────────────────────────────────────────────
    def _clear(self):
        for w in self.content.winfo_children():
            w.destroy()

    def _header(self, title, subtitle=""):
        f = tk.Frame(self.content, bg=BG)
        f.pack(fill="x", padx=30, pady=(25, 5))
        tk.Label(f, text=title, font=FONT_H1, bg=BG, fg=TEXT).pack(anchor="w")
        if subtitle:
            tk.Label(f, text=subtitle, font=FONT_BODY, bg=BG, fg=GRAY).pack(anchor="w")
        ttk.Separator(self.content, orient="horizontal").pack(fill="x", padx=30, pady=10)

    # ── Проверка первого запуска ───────────────────────────────────────────
    def _check_first_run(self):
        if not os.path.exists(OWNER_FILE):
            self.after(300, lambda: self._show_profiles(first_run=True))
        else:
            profiles = get_profiles()
            if profiles:
                self.active_profile_file.set(profiles[0][0])
                self.active_profile_data = profiles[0][1]
                self._update_profile_btn()

    def _update_profile_btn(self):
        data = self.active_profile_data
        if data.get("_type") == "owner":
            label = "👤 Я сама"
        else:
            label = f"🏢 {data.get('name', 'Клиент')}"
        self.profile_btn.config(text=label)

    # ══════════════════════════════════════════════════════════════════════
    #  Главная страница
    # ══════════════════════════════════════════════════════════════════════
    def _show_home(self):
        self._clear()
        self._header("Добро пожаловать!", "Выбери действие в левом меню")

        cards = [
            ("✍️", "Написать пост",       "Создай пост для соцсетей\nв голосе эксперта",  ACCENT,  self._show_generator),
            ("🔍", "НейроАудит",          "Word-отчёт по бизнесу\nклиента за 5 минут",    ACCENT2, self._show_audit),
            ("📅", "Контент-план",        "Темы на 2 недели\nв Excel файле",               ORANGE,  self._show_calendar),
            ("👥", "Профили",             "Добавить клиента\nили изменить свой профиль",   GRAY,    self._show_profiles),
        ]

        grid = tk.Frame(self.content, bg=BG)
        grid.pack(padx=30, pady=10)

        for i, (icon, title, desc, color, cmd) in enumerate(cards):
            card = tk.Frame(grid, bg=WHITE, relief="flat", bd=0,
                            highlightbackground=LIGHT, highlightthickness=1)
            card.grid(row=i//2, column=i%2, padx=12, pady=12, sticky="nsew")

            inner = tk.Frame(card, bg=WHITE)
            inner.pack(padx=25, pady=20)

            tk.Label(inner, text=icon, font=("Arial", 32), bg=WHITE).pack()
            tk.Label(inner, text=title, font=FONT_H2, bg=WHITE, fg=TEXT).pack(pady=(5,2))
            tk.Label(inner, text=desc, font=FONT_BODY, bg=WHITE, fg=GRAY,
                     justify="center").pack()

            btn = tk.Button(inner, text="Открыть →", font=FONT_BTN,
                            bg=color, fg=WHITE, relief="flat",
                            padx=20, pady=8, cursor="hand2", command=cmd)
            btn.pack(pady=(15,0))

    # ══════════════════════════════════════════════════════════════════════
    #  Выбор профиля
    # ══════════════════════════════════════════════════════════════════════
    def _open_profile_selector(self):
        profiles = get_profiles()
        if not profiles:
            messagebox.showinfo("Профили", "Профилей пока нет. Создай свой в разделе «Профили».")
            return

        win = tk.Toplevel(self)
        win.title("Выбрать профиль")
        win.geometry("400x350")
        win.configure(bg=BG)
        win.resizable(False, False)
        win.grab_set()

        tk.Label(win, text="Выбери профиль", font=FONT_H2, bg=BG, fg=TEXT).pack(pady=20)

        listbox = tk.Listbox(win, font=FONT_BODY, relief="flat",
                             bg=WHITE, fg=TEXT, selectbackground=ACCENT,
                             selectforeground=WHITE, bd=1, height=10)
        listbox.pack(fill="both", padx=20, expand=True)

        for fname, data in profiles:
            if data.get("_type") == "owner":
                label = "👤 Я сама (мой личный бренд)"
            else:
                label = f"🏢 {data.get('name', fname)}"
            listbox.insert("end", label)

        def select():
            sel = listbox.curselection()
            if not sel:
                return
            fname, data = profiles[sel[0]]
            self.active_profile_file.set(fname)
            self.active_profile_data = data
            self._update_profile_btn()
            win.destroy()

        tk.Button(win, text="Выбрать", font=FONT_BTN,
                  bg=ACCENT, fg=WHITE, relief="flat",
                  padx=20, pady=8, cursor="hand2", command=select).pack(pady=15)

    # ══════════════════════════════════════════════════════════════════════
    #  Управление профилями
    # ══════════════════════════════════════════════════════════════════════
    def _show_profiles(self, first_run=False):
        self._clear()
        self._header("Профили", "Твой бренд и профили клиентов")

        if first_run:
            tk.Label(self.content,
                     text="👋  Привет! Сначала создадим твой профиль эксперта.\n"
                          "Это займёт 5 минут и делается один раз.",
                     font=FONT_BODY, bg=BG, fg=TEXT).pack(pady=5)

        # Список профилей
        frame = tk.Frame(self.content, bg=BG)
        frame.pack(fill="both", expand=True, padx=30)

        profiles = get_profiles()
        if profiles:
            tk.Label(frame, text="Существующие профили:", font=("Arial",11,"bold"),
                     bg=BG, fg=TEXT).pack(anchor="w", pady=(0,8))

            for fname, data in profiles:
                row = tk.Frame(frame, bg=WHITE, highlightbackground=LIGHT,
                               highlightthickness=1)
                row.pack(fill="x", pady=3)
                icon = "👤" if data.get("_type") == "owner" else "🏢"
                name = data.get("name") or data.get("who_i_am","")[:45]
                tk.Label(row, text=f"{icon}  {name}", font=FONT_BODY,
                         bg=WHITE, fg=TEXT).pack(side="left", padx=15, pady=10)

        tk.Label(frame, text="", bg=BG).pack(pady=5)

        # Кнопки
        btns = tk.Frame(frame, bg=BG)
        btns.pack(anchor="w")

        if not os.path.exists(OWNER_FILE):
            tk.Button(btns, text="✍️  Создать МОЙ профиль (эксперт)",
                      font=FONT_BTN, bg=ACCENT, fg=WHITE, relief="flat",
                      padx=20, pady=10, cursor="hand2",
                      command=lambda: self._open_profile_form("owner")).pack(side="left", padx=(0,10))
        else:
            tk.Button(btns, text="✏️  Изменить мой профиль",
                      font=FONT_BTN, bg=GRAY, fg=WHITE, relief="flat",
                      padx=20, pady=10, cursor="hand2",
                      command=lambda: self._open_profile_form("owner")).pack(side="left", padx=(0,10))

        tk.Button(btns, text="➕  Добавить клиента",
                  font=FONT_BTN, bg=ACCENT2, fg=WHITE, relief="flat",
                  padx=20, pady=10, cursor="hand2",
                  command=lambda: self._open_profile_form("client")).pack(side="left")

    def _open_profile_form(self, ptype):
        is_owner = (ptype == "owner")
        title    = "Мой профиль (Эксперт)" if is_owner else "Новый клиент"

        questions = [
            ("name",            "Имя / название бренда"),
            ("who_i_am",        "Чем занимаешься? (1-2 предложения)"),
            ("target_client",   "Кому помогаешь? (описание клиента и его боль)"),
            ("main_result",     "Главный результат от работы с тобой"),
            ("differentiation", "Чем отличаешься от других?"),
            ("tone",            "Стиль общения (тепло / по-деловому / с юмором / строго)"),
            ("my_phrases",      "3 фразы которые ты реально используешь (через запятую)"),
            ("never_say",       "Что НИКОГДА не скажешь клиенту (через запятую)"),
            ("offer",           "Оффер одним предложением (что + для кого + результат)"),
            ("price_base",      "Цена базовой услуги (только число, например 12000)"),
        ]
        if is_owner:
            questions.insert(4, ("archetypes", "Твои архетипы (Герой, Правитель, Маг — через запятую)"))

        # Загрузить существующие данные если есть
        existing = {}
        if is_owner and os.path.exists(OWNER_FILE):
            existing = load_json(OWNER_FILE)

        win = tk.Toplevel(self)
        win.title(title)
        win.geometry("620x700")
        win.configure(bg=BG)
        win.grab_set()

        tk.Label(win, text=title, font=FONT_H1, bg=BG, fg=TEXT).pack(pady=(20,5), padx=30, anchor="w")
        tk.Label(win, text="Заполни поля — это Паспорт Бренда", font=FONT_BODY,
                 bg=BG, fg=GRAY).pack(padx=30, anchor="w")

        canvas = tk.Canvas(win, bg=BG, highlightthickness=0)
        scroll = ttk.Scrollbar(win, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True, padx=(20,0), pady=10)

        inner = tk.Frame(canvas, bg=BG)
        canvas_win = canvas.create_window((0,0), window=inner, anchor="nw")

        def on_resize(e):
            canvas.itemconfig(canvas_win, width=e.width)
        canvas.bind("<Configure>", on_resize)

        def on_scroll(e):
            canvas.yview_scroll(int(-1*(e.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", on_scroll)

        inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        entries = {}
        for key, label in questions:
            tk.Label(inner, text=label, font=("Arial",10,"bold"),
                     bg=BG, fg=TEXT).pack(anchor="w", padx=10, pady=(10,2))
            is_long = key in ("who_i_am","target_client","main_result","differentiation","offer")
            if is_long:
                w = tk.Text(inner, font=FONT_BODY, height=3, relief="solid",
                            bd=1, wrap="word", bg=WHITE)
                w.pack(fill="x", padx=10, ipady=4)
                val = existing.get(key, "")
                if isinstance(val, list):
                    val = ", ".join(val)
                if val:
                    w.insert("1.0", val)
            else:
                w = tk.Entry(inner, font=FONT_BODY, relief="solid", bd=1, bg=WHITE)
                w.pack(fill="x", padx=10, ipady=6)
                val = existing.get(key, "")
                if isinstance(val, list):
                    val = ", ".join(val)
                if val:
                    w.insert(0, str(val))
            entries[key] = (w, is_long)

        def save():
            data = {"_type": "owner" if is_owner else "client"}
            for key, (widget, is_long) in entries.items():
                if is_long:
                    val = widget.get("1.0", "end-1c").strip()
                else:
                    val = widget.get().strip()
                if not val and key == "name":
                    messagebox.showerror("Ошибка", "Заполни поле «Имя / название бренда»")
                    return
                if key in ("my_phrases", "never_say", "archetypes"):
                    data[key] = [p.strip() for p in val.split(",") if p.strip()]
                elif key == "price_base":
                    try:
                        data[key] = int(val.replace(" ","").replace("₽",""))
                    except ValueError:
                        data[key] = val
                else:
                    data[key] = val

            if is_owner:
                save_json(OWNER_FILE, data)
                self.active_profile_file.set("owner.json")
                self.active_profile_data = data
                self._update_profile_btn()
            else:
                fname = name_to_filename(data.get("name","client")) + ".json"
                path  = os.path.join(PROFILES_DIR, fname)
                counter = 1
                base = path.replace(".json","")
                while os.path.exists(path):
                    path = f"{base}_{counter}.json"
                    counter += 1
                save_json(path, data)

            win.destroy()
            messagebox.showinfo("Готово!", f"Профиль «{data.get('name','')}» сохранён ✅")
            self._show_profiles()

        tk.Button(inner, text="💾  Сохранить профиль", font=FONT_BTN,
                  bg=ACCENT, fg=WHITE, relief="flat",
                  padx=25, pady=12, cursor="hand2", command=save).pack(pady=20)

    # ══════════════════════════════════════════════════════════════════════
    #  Генератор постов
    # ══════════════════════════════════════════════════════════════════════
    def _show_generator(self):
        self._clear()
        brand = self.active_profile_data.get("name") or "не выбран"
        self._header("✍️  Написать пост", f"Профиль: {brand}")

        if not self.active_profile_data:
            self._no_profile_msg()
            return

        form = tk.Frame(self.content, bg=BG)
        form.pack(fill="both", expand=True, padx=30)

        # Тема
        tk.Label(form, text="Тема поста:", font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w")
        topic_entry = tk.Entry(form, font=FONT_BODY, relief="solid", bd=1, bg=WHITE)
        topic_entry.pack(fill="x", ipady=8, pady=(3,15))

        row = tk.Frame(form, bg=BG)
        row.pack(fill="x")

        # Тип поста
        tk.Label(row, text="Тип поста:", font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w")
        post_type = tk.StringVar(value="полезный")
        for val, label in [("полезный","📚 Полезный"),("история","📖 История"),
                           ("кейс","🏆 Кейс"),("анонс","📣 Анонс")]:
            tk.Radiobutton(row, text=label, variable=post_type, value=val,
                           font=FONT_BODY, bg=BG, fg=TEXT,
                           activebackground=BG).pack(anchor="w")

        # Платформа
        tk.Label(form, text="\nПлатформа:", font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w")
        platform = tk.StringVar(value="TG")
        prow = tk.Frame(form, bg=BG)
        prow.pack(anchor="w")
        for val, label in [("TG","Telegram"),("VK","ВКонтакте"),("MAX","MAX"),("все","Все сразу")]:
            tk.Radiobutton(prow, text=label, variable=platform, value=val,
                           font=FONT_BODY, bg=BG, fg=TEXT,
                           activebackground=BG).pack(side="left", padx=10)

        # Кнопка
        gen_btn = tk.Button(form, text="⚡  Сгенерировать пост", font=FONT_BTN,
                            bg=ACCENT, fg=WHITE, relief="flat",
                            padx=25, pady=12, cursor="hand2")
        gen_btn.pack(pady=15, anchor="w")

        # Статус
        status_lbl = tk.Label(form, text="", font=FONT_BODY, bg=BG, fg=GRAY)
        status_lbl.pack(anchor="w")

        # Результат
        tk.Label(form, text="Результат:", font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w", pady=(10,3))
        result_box = scrolledtext.ScrolledText(form, font=("Arial",10), height=12,
                                               relief="solid", bd=1, wrap="word", bg=WHITE)
        result_box.pack(fill="both", expand=True)

        save_btn = tk.Button(form, text="💾  Сохранить в файл", font=FONT_BTN,
                             bg=ACCENT2, fg=WHITE, relief="flat",
                             padx=20, pady=8, cursor="hand2", state="disabled")
        save_btn.pack(pady=10, anchor="w")

        def do_generate():
            topic = topic_entry.get().strip()
            if not topic:
                messagebox.showwarning("Внимание", "Введи тему поста")
                return

            gen_btn.config(state="disabled", text="⏳ Генерирую...")
            status_lbl.config(text="Подключаюсь к Gemini AI...", fg=ORANGE)
            result_box.delete("1.0", "end")
            save_btn.config(state="disabled")

            def task():
                try:
                    from dotenv import load_dotenv
                    import google.generativeai as genai
                    import os as _os
                    load_dotenv(os.path.join(BASE_DIR, ".env"))
                    api_key = _os.getenv("GEMINI_API_KEY")
                    if not api_key:
                        self.after(0, lambda: status_lbl.config(
                            text="❌ GEMINI_API_KEY не найден в .env файле", fg=RED))
                        return

                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel("gemini-1.5-flash")

                    profile = self.active_profile_data
                    phrases   = ", ".join(profile.get("my_phrases", []))
                    never_say = ", ".join(profile.get("never_say", []))
                    p_name    = profile.get("name") or "эксперт"

                    plat_instr = {
                        "TG":  "Формат Telegram: до 4096 символов, можно эмодзи.",
                        "VK":  "Формат ВКонтакте: добавь 5-7 хэштегов в конце.",
                        "MAX": "Формат MAX: дружелюбный тон, эмодзи умеренно.",
                        "все": "Напиши версию для Telegram. Затем напиши версию для ВКонтакте с хэштегами.",
                    }.get(platform.get(), "")

                    type_instr = {
                        "полезный": "Структура: заголовок → проблема → 3-5 советов → вывод → призыв.",
                        "история":  "Структура: яркое начало → развитие → кульминация → вывод.",
                        "кейс":     "Структура: ДО → что сделали → результат в цифрах → вывод.",
                        "анонс":    "Структура: интрига → ценность → призыв к действию.",
                    }.get(post_type.get(), "")

                    prompt = f"""Ты копирайтер эксперта. Паспорт Бренда:
Имя: {p_name}
Кто я: {profile.get('who_i_am','')}
Клиент: {profile.get('target_client','')}
Тон: {profile.get('tone','')}
Мои фразы: {phrases}
Никогда не говорю: {never_say}
Оффер: {profile.get('offer','')}

Напиши пост от первого лица в голосе эксперта.
Тема: {topic}
{type_instr}
{plat_instr}

После текста добавь строку:
ПРОМПТ ДЛЯ ВИЗУАЛА: [описание картинки для нейросети]"""

                    response = model.generate_content(prompt)
                    text = response.text

                    # Сохраняем файл
                    os.makedirs(OUTPUT_POSTS, exist_ok=True)
                    today    = date.today().strftime("%Y-%m-%d")
                    plat_tag = platform.get()
                    fpath    = os.path.join(OUTPUT_POSTS, f"{today}_{plat_tag}.txt")
                    with open(fpath, "w", encoding="utf-8") as fout:
                        fout.write(f"Тема: {topic}\nДата: {today}\n\n{'─'*50}\n\n{text}")

                    def update():
                        result_box.delete("1.0","end")
                        result_box.insert("1.0", text)
                        status_lbl.config(text=f"✅ Готово! Сохранено: {fpath}", fg=ACCENT)
                        gen_btn.config(state="normal", text="⚡  Сгенерировать пост")
                        save_btn.config(state="normal")
                    self.after(0, update)

                except Exception as e:
                    self.after(0, lambda: [
                        status_lbl.config(text=f"❌ Ошибка: {e}", fg=RED),
                        gen_btn.config(state="normal", text="⚡  Сгенерировать пост"),
                    ])

            run_in_thread(task)

        gen_btn.config(command=do_generate)

        def do_save():
            text = result_box.get("1.0","end-1c")
            if not text.strip():
                return
            fpath = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Text files","*.txt")],
                initialdir=OUTPUT_POSTS,
                title="Сохранить пост",
            )
            if fpath:
                with open(fpath,"w",encoding="utf-8") as f:
                    f.write(text)
                messagebox.showinfo("Готово", f"Сохранено: {fpath}")

        save_btn.config(command=do_save)

    # ══════════════════════════════════════════════════════════════════════
    #  НейроАудит
    # ══════════════════════════════════════════════════════════════════════
    def _show_audit(self):
        self._clear()
        self._header("🔍  НейроАудит клиента",
                     "Анализ бизнеса клиента → Word-отчёт за 5 минут")

        if not os.path.exists(OWNER_FILE):
            self._no_profile_msg()
            return

        owner = load_json(OWNER_FILE)
        expert_name = owner.get("name") or "Эксперт"

        form = tk.Frame(self.content, bg=BG)
        form.pack(fill="both", expand=True, padx=30)

        tk.Label(form, text=f"Эксперт (ты): {expert_name}",
                 font=("Arial",10), bg=BG, fg=ACCENT).pack(anchor="w", pady=(0,10))

        fields = [
            ("client_name", "Имя/название бизнеса клиента *", False),
            ("niche",       "Ниша клиента *", False),
            ("site_text",   "Текст сайта или оффера клиента * (скопируй и вставь)", True),
            ("audience",    "Целевая аудитория (необязательно)", False),
        ]

        entries = {}
        for key, label, multiline in fields:
            tk.Label(form, text=label, font=("Arial",10,"bold"), bg=BG, fg=TEXT).pack(anchor="w", pady=(8,2))
            if multiline:
                w = tk.Text(form, font=FONT_BODY, height=6, relief="solid", bd=1, wrap="word", bg=WHITE)
                w.pack(fill="x", ipady=4)
            else:
                w = tk.Entry(form, font=FONT_BODY, relief="solid", bd=1, bg=WHITE)
                w.pack(fill="x", ipady=8)
            entries[key] = (w, multiline)

        status_lbl = tk.Label(form, text="", font=FONT_BODY, bg=BG, fg=GRAY)
        status_lbl.pack(anchor="w", pady=5)

        audit_btn = tk.Button(form, text="🔍  Сделать НейроАудит", font=FONT_BTN,
                              bg=ACCENT2, fg=WHITE, relief="flat",
                              padx=25, pady=12, cursor="hand2")
        audit_btn.pack(anchor="w")

        def do_audit():
            vals = {}
            for key, (widget, ml) in entries.items():
                vals[key] = widget.get("1.0","end-1c").strip() if ml else widget.get().strip()

            if not vals["client_name"] or not vals["niche"] or not vals["site_text"]:
                messagebox.showwarning("Внимание", "Заполни обязательные поля (отмечены *)")
                return

            audit_btn.config(state="disabled", text="⏳ Анализирую...")
            status_lbl.config(text="Подключаюсь к Gemini AI... (~30 секунд)", fg=ORANGE)

            def task():
                try:
                    from dotenv import load_dotenv
                    import google.generativeai as genai
                    import os as _os
                    load_dotenv(os.path.join(BASE_DIR, ".env"))
                    api_key = _os.getenv("GEMINI_API_KEY")
                    if not api_key:
                        self.after(0, lambda: status_lbl.config(
                            text="❌ GEMINI_API_KEY не найден в .env", fg=RED))
                        return

                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel("gemini-1.5-flash")

                    prompt = f"""Ты — ИИ-Стратег, эксперт по маркетингу. Аудит проводит: {expert_name}

Проведи честный анализ бизнеса клиента.

Клиент: {vals['client_name']}
Ниша: {vals['niche']}
{"Аудитория: " + vals['audience'] if vals['audience'] else ""}

Текст сайта / оффера:
{vals['site_text']}

Напиши аудит строго в трёх разделах:

## РЕНТГЕН БИЗНЕСА
3 главные дыры. Для каждой: в чём проблема, почему теряет деньги, приоритет.

## ДОРОЖНАЯ КАРТА
4 конкретных шага: что сделать, срок, результат.

## СКРИПТ ЗАКРЫТИЯ
Как {expert_name} предложит свои услуги после аудита. 2-3 конкретные фразы."""

                    response = model.generate_content(prompt)
                    audit_text = response.text

                    # Word отчёт
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    os.makedirs(OUTPUT_AUDIT, exist_ok=True)
                    today  = date.today().strftime("%Y-%m-%d")
                    safe   = vals['client_name'].replace(" ","_")
                    fpath  = os.path.join(OUTPUT_AUDIT, f"НейроАудит_{safe}_{today}.docx")

                    doc = Document()
                    doc.styles["Normal"].font.name = "Arial"
                    doc.styles["Normal"].font.size = Pt(11)

                    t = doc.add_heading("НейроАудит", level=0)
                    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    t.runs[0].font.color.rgb = RGBColor(0x2C,0x3E,0x50)

                    s = doc.add_paragraph(f"{vals['client_name']} • {vals['niche']} • {today}")
                    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    s.runs[0].font.color.rgb = RGBColor(0x7F,0x8C,0x8D)

                    b = doc.add_paragraph(f"Подготовлено: {expert_name}")
                    b.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    b.runs[0].font.color.rgb = RGBColor(0x27,0xAE,0x60)
                    doc.add_paragraph()

                    for section in audit_text.split("##"):
                        section = section.strip()
                        if not section:
                            continue
                        lines = section.split("\n",1)
                        sec_title = lines[0].strip()
                        sec_body  = lines[1].strip() if len(lines)>1 else ""
                        if sec_title:
                            h = doc.add_heading(sec_title, level=1)
                            h.runs[0].font.color.rgb = RGBColor(0x27,0xAE,0x60)
                        for para in sec_body.split("\n\n"):
                            para = para.strip()
                            if para:
                                doc.add_paragraph(para).paragraph_format.space_after = Pt(6)
                        doc.add_paragraph()

                    doc.save(fpath)

                    def done():
                        status_lbl.config(
                            text=f"✅ Готово! Открываю файл...", fg=ACCENT)
                        audit_btn.config(state="normal", text="🔍  Сделать НейроАудит")
                        os.startfile(fpath)
                        messagebox.showinfo("НейроАудит готов! 🎉",
                            f"Word-файл сохранён:\n{fpath}\n\nФайл открывается автоматически.")
                    self.after(0, done)

                except Exception as e:
                    self.after(0, lambda: [
                        status_lbl.config(text=f"❌ Ошибка: {e}", fg=RED),
                        audit_btn.config(state="normal", text="🔍  Сделать НейроАудит"),
                    ])

            run_in_thread(task)

        audit_btn.config(command=do_audit)

    # ══════════════════════════════════════════════════════════════════════
    #  Контент-план
    # ══════════════════════════════════════════════════════════════════════
    def _show_calendar(self):
        self._clear()
        brand = self.active_profile_data.get("name") or "не выбран"
        self._header("📅  Контент-план", f"Профиль: {brand}")

        if not self.active_profile_data:
            self._no_profile_msg()
            return

        form = tk.Frame(self.content, bg=BG)
        form.pack(padx=30, fill="x")

        tk.Label(form, text="Постов в неделю:", font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w", pady=(10,3))
        ppw_var = tk.IntVar(value=3)
        ppw_row = tk.Frame(form, bg=BG)
        ppw_row.pack(anchor="w")
        for n in range(1,8):
            tk.Radiobutton(ppw_row, text=str(n), variable=ppw_var, value=n,
                           font=FONT_BODY, bg=BG, activebackground=BG).pack(side="left")

        tk.Label(form, text="\nДата старта (ДД.ММ.ГГГГ, пусто = сегодня):",
                 font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w", pady=(5,3))
        date_entry = tk.Entry(form, font=FONT_BODY, relief="solid", bd=1, bg=WHITE, width=20)
        date_entry.pack(anchor="w", ipady=8)

        status_lbl = tk.Label(form, text="", font=FONT_BODY, bg=BG, fg=GRAY)
        status_lbl.pack(anchor="w", pady=8)

        cal_btn = tk.Button(form, text="📅  Создать контент-план", font=FONT_BTN,
                            bg=ORANGE, fg=WHITE, relief="flat",
                            padx=25, pady=12, cursor="hand2")
        cal_btn.pack(anchor="w")

        def do_calendar():
            from datetime import datetime
            start_str = date_entry.get().strip()
            if start_str:
                try:
                    start_date = datetime.strptime(start_str, "%d.%m.%Y").date()
                except ValueError:
                    messagebox.showerror("Ошибка", "Формат даты: ДД.ММ.ГГГГ")
                    return
            else:
                start_date = date.today()

            cal_btn.config(state="disabled", text="⏳ Генерирую темы...")
            status_lbl.config(text="Подключаюсь к Gemini AI...", fg=ORANGE)

            def task():
                try:
                    from dotenv import load_dotenv
                    import google.generativeai as genai
                    import os as _os
                    load_dotenv(os.path.join(BASE_DIR, ".env"))
                    api_key = _os.getenv("GEMINI_API_KEY")
                    if not api_key:
                        self.after(0, lambda: status_lbl.config(text="❌ GEMINI_API_KEY не найден", fg=RED))
                        return

                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel("gemini-1.5-flash")
                    profile = self.active_profile_data
                    ppw     = ppw_var.get()
                    total   = ppw * 2

                    prompt = f"""Ты контент-стратег. Создай план для:
Бренд: {profile.get('name','')}
Ниша: {profile.get('who_i_am','')}
Тон: {profile.get('tone','')}

Придумай {total} тем на 2 недели. Чередуй: полезный, история, кейс, анонс, вирусный, личное.
Формат (одна тема на строку):
ТИП|ТЕМА|ХЭШТЕГИ

Ровно {total} строк, без лишнего текста."""

                    response = model.generate_content(prompt)
                    raw = response.text.strip()

                    topics = []
                    for line in raw.split("\n"):
                        if "|" in line:
                            parts = line.split("|",2)
                            if len(parts) >= 2:
                                topics.append((parts[0].strip(), parts[1].strip(),
                                               parts[2].strip() if len(parts)>2 else ""))
                        if len(topics) >= total:
                            break

                    # Excel
                    import openpyxl
                    from openpyxl.styles import Font as XFont, PatternFill, Alignment, Border, Side
                    from openpyxl.utils import get_column_letter
                    from datetime import timedelta

                    DAYS = ["Понедельник","Вторник","Среда","Четверг","Пятница","Суббота","Воскресенье"]
                    os.makedirs(OUTPUT_DIR, exist_ok=True)
                    today_str = date.today().strftime("%Y-%m-%d")
                    xlsx_path = os.path.join(BASE_DIR,"output",f"calendar_{today_str}.xlsx")

                    wb = openpyxl.Workbook()
                    ws = wb.active
                    ws.title = "Контент-план"

                    border = Border(left=Side(style="thin"),right=Side(style="thin"),
                                    top=Side(style="thin"),bottom=Side(style="thin"))
                    hfill = PatternFill(start_color="2C3E50",end_color="2C3E50",fill_type="solid")
                    hfont = XFont(name="Arial",bold=True,color="FFFFFF",size=11)

                    for col, h in enumerate(["Дата","День","Тип","Тема","Хэштеги","Статус"],1):
                        c = ws.cell(row=1,column=col,value=h)
                        c.fill,c.font = hfill,hfont
                        c.alignment = Alignment(horizontal="center")
                        c.border = border
                    ws.row_dimensions[1].height = 25

                    step    = max(1, 7//ppw)
                    offsets = [i*step for i in range(ppw)]
                    idx, row = 0, 2
                    for week in range(2):
                        wstart = start_date + timedelta(weeks=week)
                        for off in offsets:
                            if idx >= len(topics):
                                break
                            pdate = wstart + timedelta(days=off)
                            ptype,theme,tags = topics[idx]
                            fill = PatternFill(start_color="ECF0F1",end_color="ECF0F1",fill_type="solid") if row%2==0 \
                                   else PatternFill(start_color="FFFFFF",end_color="FFFFFF",fill_type="solid")
                            for col,val in enumerate([pdate.strftime("%d.%m.%Y"),
                                DAYS[pdate.weekday()],ptype,theme,tags,"📝 Не начат"],1):
                                c = ws.cell(row=row,column=col,value=val)
                                c.border = border
                                c.alignment = Alignment(vertical="center",wrap_text=True)
                                c.fill = fill
                                c.font = XFont(name="Arial",size=10)
                            ws.row_dimensions[row].height = 40
                            idx += 1
                            row += 1

                    for i,w in enumerate([14,14,16,45,35,14],1):
                        ws.column_dimensions[get_column_letter(i)].width = w

                    wb.save(xlsx_path)

                    def done():
                        status_lbl.config(text=f"✅ Готово! {len(topics)} постов.", fg=ACCENT)
                        cal_btn.config(state="normal", text="📅  Создать контент-план")
                        os.startfile(xlsx_path)
                    self.after(0, done)

                except Exception as e:
                    self.after(0, lambda: [
                        status_lbl.config(text=f"❌ Ошибка: {e}", fg=RED),
                        cal_btn.config(state="normal", text="📅  Создать контент-план"),
                    ])

            run_in_thread(task)

        cal_btn.config(command=do_calendar)

    # ══════════════════════════════════════════════════════════════════════
    #  Публикация TG / VK
    # ══════════════════════════════════════════════════════════════════════
    def _show_publish_tg(self):
        self._show_publish("TG")

    def _show_publish_vk(self):
        self._show_publish("VK")

    def _show_publish(self, platform):
        self._clear()
        self._header(f"📢  Публикация в {platform}",
                     "Выбери файл поста и нажми «Опубликовать»")

        form = tk.Frame(self.content, bg=BG)
        form.pack(padx=30, fill="x")

        tk.Label(form, text="Файл поста:", font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w", pady=(10,3))
        file_var = tk.StringVar()
        file_row = tk.Frame(form, bg=BG)
        file_row.pack(fill="x")
        tk.Entry(file_row, textvariable=file_var, font=FONT_BODY,
                 relief="solid", bd=1, bg=WHITE).pack(side="left", fill="x", expand=True, ipady=8)

        def browse():
            f = filedialog.askopenfilename(initialdir=OUTPUT_POSTS,
                filetypes=[("Text","*.txt"),("All","*.*")])
            if f:
                file_var.set(f)
        tk.Button(file_row, text="Выбрать", font=FONT_BTN,
                  bg=GRAY, fg=WHITE, relief="flat", padx=12, pady=6,
                  cursor="hand2", command=browse).pack(side="left", padx=(8,0))

        tk.Label(form, text="\nКогда публиковать:", font=("Arial",11,"bold"), bg=BG, fg=TEXT).pack(anchor="w")
        when_var = tk.StringVar(value="now")
        tk.Radiobutton(form, text="Прямо сейчас", variable=when_var, value="now",
                       font=FONT_BODY, bg=BG, activebackground=BG).pack(anchor="w")
        sched_row = tk.Frame(form, bg=BG)
        sched_row.pack(anchor="w")
        tk.Radiobutton(sched_row, text="В ", variable=when_var, value="scheduled",
                       font=FONT_BODY, bg=BG, activebackground=BG).pack(side="left")
        time_entry = tk.Entry(sched_row, font=FONT_BODY, relief="solid", bd=1,
                              bg=WHITE, width=8)
        time_entry.insert(0, "10:00")
        time_entry.pack(side="left", ipady=6)
        tk.Label(sched_row, text="(ЧЧ:ММ)", font=FONT_BODY, bg=BG, fg=GRAY).pack(side="left", padx=5)

        status_lbl = tk.Label(form, text="", font=FONT_BODY, bg=BG, fg=GRAY)
        status_lbl.pack(anchor="w", pady=8)

        pub_btn = tk.Button(form, text=f"📢  Опубликовать в {platform}", font=FONT_BTN,
                            bg=ACCENT, fg=WHITE, relief="flat",
                            padx=25, pady=12, cursor="hand2")
        pub_btn.pack(anchor="w")

        def do_publish():
            fpath = file_var.get().strip()
            if not fpath or not os.path.exists(fpath):
                messagebox.showwarning("Внимание", "Выбери файл поста")
                return

            script = "publisher_tg.py" if platform == "TG" else "publisher_vk.py"
            if when_var.get() == "now":
                args = ["--file", fpath, "--now"]
            else:
                t = time_entry.get().strip()
                args = ["--file", fpath, "--time", t]

            pub_btn.config(state="disabled", text="⏳ Публикую...")
            status_lbl.config(text=f"Отправляю в {platform}...", fg=ORANGE)

            def task():
                try:
                    result = subprocess.run(
                        [sys.executable, os.path.join(BASE_DIR, script)] + args,
                        capture_output=True, text=True, cwd=BASE_DIR
                    )
                    out = result.stdout + result.stderr
                    if "✅" in out or "опубликован" in out.lower():
                        self.after(0, lambda: [
                            status_lbl.config(text=f"✅ Пост опубликован в {platform}!", fg=ACCENT),
                            pub_btn.config(state="normal", text=f"📢  Опубликовать в {platform}"),
                        ])
                    else:
                        self.after(0, lambda: [
                            status_lbl.config(text=f"⚠️ {out[:120]}", fg=ORANGE),
                            pub_btn.config(state="normal", text=f"📢  Опубликовать в {platform}"),
                        ])
                except Exception as e:
                    self.after(0, lambda: [
                        status_lbl.config(text=f"❌ Ошибка: {e}", fg=RED),
                        pub_btn.config(state="normal", text=f"📢  Опубликовать в {platform}"),
                    ])

            run_in_thread(task)

        pub_btn.config(command=do_publish)

    # ── Заглушка для отсутствующего профиля ───────────────────────────────
    def _no_profile_msg(self):
        tk.Label(self.content,
                 text="⚠️  Сначала создай профиль в разделе «Профили»",
                 font=FONT_H2, bg=BG, fg=ORANGE).pack(pady=40)
        tk.Button(self.content, text="👥  Перейти к профилям",
                  font=FONT_BTN, bg=ACCENT2, fg=WHITE, relief="flat",
                  padx=20, pady=10, cursor="hand2",
                  command=self._show_profiles).pack()


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    os.chdir(BASE_DIR)
    app = App()
    app.mainloop()
