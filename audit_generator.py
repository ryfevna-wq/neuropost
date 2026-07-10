"""
НейроПост v2.0 — Модуль 5: Генератор НейроАудита (Режим А)

owner  = эксперт (ты) — подписывает отчёт, даёт рекомендации
client = клиент      — чей бизнес анализируем
"""
import argparse
import json
import os
import sys
from datetime import date
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

OUTPUT_DIR  = "output/audits"
PROMPT_FILE = "prompts/audit_base.txt"


def load_profile(path):
    if not os.path.exists(path):
        print(f"\n❌ Профиль не найден: {path}\n")
        sys.exit(1)
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def load_audit_prompt():
    if os.path.exists(PROMPT_FILE):
        with open(PROMPT_FILE, encoding="utf-8") as f:
            return f.read().strip()
    return None


def build_system_prompt(owner):
    base = load_audit_prompt()
    if base:
        return base

    expert_name = owner.get("name") or owner.get("who_i_am", "Эксперт")[:40]
    return f"""Ты — ИИ-Стратег, эксперт по маркетингу и продажам.
Аудит проводит: {expert_name}

Твоя задача — честно проанализировать бизнес клиента и найти ключевые точки роста.

Структура аудита строго в трёх разделах:

## РЕНТГЕН БИЗНЕСА
3 главные дыры. Для каждой:
- В чём проблема (конкретно)
- Почему это теряет деньги/клиентов
- Приоритет: высокий / средний / низкий

## ДОРОЖНАЯ КАРТА
4 конкретных шага. Каждый:
- Что именно сделать
- Срок (неделя 1, неделя 2-4 и т.д.)
- Ожидаемый результат

## СКРИПТ ЗАКРЫТИЯ
Как {expert_name} может предложить свои услуги после аудита.
2-3 конкретные фразы: реплика → реакция клиента → ответ на возражение.

Пиши честно и конкретно. Каждое утверждение обоснованно."""


def run_audit(owner, client_name, niche, site_text, audience):
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("\n❌ GEMINI_API_KEY не найден в .env файле\n")
        sys.exit(1)

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")

    audience_line = f"\nЦелевая аудитория: {audience}" if audience else ""

    prompt = f"""{build_system_prompt(owner)}

Проведи НейроАудит:

Клиент: {client_name}
Ниша: {niche}{audience_line}

Текст сайта / оффера:
{site_text}"""

    print("\n⏳ Анализирую бизнес клиента...")
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"\n❌ Ошибка: {e}\n")
        sys.exit(1)


def create_word_report(audit_text, client_name, niche, owner):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    today = date.today().strftime("%Y-%m-%d")
    safe  = client_name.replace(" ", "_").replace("/", "-")
    filepath = os.path.join(OUTPUT_DIR, f"НейроАудит_{safe}_{today}.docx")

    expert_name = owner.get("name") or owner.get("who_i_am", "")[:40]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("НейроАудит", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    sub = doc.add_paragraph(f"{client_name} • {niche} • {today}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    by = doc.add_paragraph(f"Подготовлено: {expert_name}")
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by.runs[0].font.size = Pt(10)
    by.runs[0].font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    doc.add_paragraph()

    for section in audit_text.split("##"):
        section = section.strip()
        if not section:
            continue
        lines = section.split("\n", 1)
        sec_title = lines[0].strip()
        sec_body  = lines[1].strip() if len(lines) > 1 else ""

        if sec_title:
            h = doc.add_heading(sec_title, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

        for para in sec_body.split("\n\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()

    footer = doc.add_paragraph("Создано с помощью НейроПост v2.0")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)

    doc.save(filepath)
    return filepath


def get_client_info_interactive():
    """Запрашивает данные клиента вручную если профиль не задан"""
    print("\nВведи данные клиента:")
    client_name = input("Имя/название бизнеса: ").strip()
    niche = input("Ниша: ").strip()
    print("\nВставь текст сайта или оффера клиента.")
    print("Пустая строка дважды = конец ввода:")
    lines = []
    while True:
        line = input()
        if line == "" and lines and lines[-1] == "":
            break
        lines.append(line)
    site_text = "\n".join(lines).strip()
    audience = input("\nЦелевая аудитория (Enter — пропустить): ").strip()
    return client_name, niche, site_text, audience


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--owner",  default="profiles/owner.json", help="Профиль эксперта")
    parser.add_argument("--client", default=None, help="Профиль клиента (json)")
    args = parser.parse_args()

    owner = load_profile(args.owner)
    expert_name = owner.get("name") or owner.get("who_i_am", "Эксперт")[:40]

    print("\n" + "─" * 50)
    print(f"   НейроАудит  |  Эксперт: {expert_name}")
    print("─" * 50)
    print("   Режим А — Базовый (3 дыры + дорожная карта + скрипт)\n")

    # Если передан профиль клиента — берём данные из него
    if args.client and os.path.exists(args.client):
        client = load_profile(args.client)
        client_name = client.get("name") or "Клиент"
        niche       = client.get("who_i_am", "")
        site_text   = (
            f"Оффер: {client.get('offer', '')}\n"
            f"Результат для клиента: {client.get('main_result', '')}\n"
            f"Отличие: {client.get('differentiation', '')}\n"
            f"Целевой клиент: {client.get('target_client', '')}"
        )
        audience = client.get("target_client", "")

        print(f"Клиент из профиля: {client_name}")
        print(f"Ниша: {niche}\n")

        # Возможность добавить живой текст сайта
        add_text = input("Добавить текст сайта/оффера клиента? (да/нет): ").strip().lower()
        if add_text in ("да", "д"):
            print("\nВставь текст (пустая строка дважды = конец):")
            lines = []
            while True:
                line = input()
                if line == "" and lines and lines[-1] == "":
                    break
                lines.append(line)
            extra = "\n".join(lines).strip()
            if extra:
                site_text = extra + "\n\n" + site_text
    else:
        # Клиент — сам владелец (аудит себя) или ввод вручную
        if args.client is None or args.client == args.owner:
            client_name, niche, site_text, audience = get_client_info_interactive()
        else:
            print(f"⚠️  Профиль не найден: {args.client}")
            client_name, niche, site_text, audience = get_client_info_interactive()

    if not site_text:
        print("❌ Нужен текст сайта или оффера")
        sys.exit(1)

    audit_text = run_audit(owner, client_name, niche, site_text, audience)
    filepath   = create_word_report(audit_text, client_name, niche, owner)

    print("\n" + "─" * 50)
    print("✅ НейроАудит готов!")
    print(f"📄 Файл: {filepath}")
    print("─" * 50)
    preview = audit_text[:600] + "..." if len(audit_text) > 600 else audit_text
    print(f"\n{preview}\n")


if __name__ == "__main__":
    main()
